"""文档摄入与解析：把各种格式的文件 / 网页解析成纯文本。

解析出的文本用于知识库预览，也是将来 RAG 检索的语料基础。
任何一种格式解析失败都抛异常，由上层记录 status=failed。

HTML 解析基于 lxml：高性能 C 实现 + XPath 正文定位 + 自动清洗。
"""
import io
from collections import deque
from urllib.parse import urldefrag, urljoin, urlparse

import httpx
from lxml import html as lxml_html
from lxml_html_clean import Cleaner

# 纯文本类：直接按 UTF-8 读取
PLAIN_EXTS = {".txt", ".md", ".markdown", ".csv", ".json", ".yaml", ".yml", ".log"}
HTML_EXTS = {".html", ".htm"}
SUPPORTED_EXTS = PLAIN_EXTS | HTML_EXTS | {".pdf", ".docx", ".xlsx"}

# 网页抓取的浏览器 UA，避免部分站点拒绝
_UA = "Mozilla/5.0 (compatible; CodeAgentKB/0.1; +knowledge-base)"

# ── XPath 正文选择器（按优先级从高到低） ──
# 优先使用语义化标签 + 常见内容选择器，尽量避开导航/侧边栏/广告/页脚
_CONTENT_SELECTORS = [
    '//article',
    '//main',
    # 文档站容器（GitBook / Docusaurus / VuePress）—— 优先于 role="main"
    '//div[contains(@class, "body-inner")]',
    '//div[contains(@class, "page-inner")]',
    '//div[contains(@class, "theme-default-content")]',
    '//*[@role="main"]',
    '//div[contains(@class, "content") or contains(@class, "post") or contains(@class, "article")]',
    '//div[@id="content" or @id="main" or @id="article"]',
    '//section[contains(@class, "content")]',
    '//div[contains(@class, "markdown") or contains(@class, "document")]',
]

# ── HTML 清洗器（单例复用） ──
_cleaner = Cleaner(
    scripts=True,
    javascript=True,
    comments=True,
    style=True,
    links=False,        # 保留 <a> 方便提取链接
    meta=True,
    processing_instructions=True,
    embedded=True,
    frames=True,
    forms=False,        # 保留表单文本
    annoying_tags=False,
    remove_unknown_tags=False,
    safe_attrs_only=False,
)


def _parse_html_doc(html_str: str) -> lxml_html.HtmlElement:
    """将 HTML 字符串解析为 lxml 文档树。"""
    return lxml_html.document_fromstring(html_str)


def _clean_html(doc: lxml_html.HtmlElement) -> lxml_html.HtmlElement:
    """清洗 HTML：去除 script/style/注释等垃圾，返回干净 DOM。"""
    return _cleaner.clean_html(doc)


# ── HTML → 结构化文本（保留标题层级） ──

# 块级标签：会在前后产生换行
_BLOCK_TAGS = {
    "p", "div", "section", "article", "main", "header", "footer",
    "nav", "aside", "form", "table", "tr", "ul", "ol", "dl",
    "blockquote", "figure", "figcaption", "details", "summary",
    "hr", "br",
}

# 需要跳过的标签（纯布局/导航/广告/脚本残留），不产生任何输出
_SKIP_TAGS = {
    "script", "style", "noscript", "nav", "footer", "aside",
    "iframe", "object", "embed", "canvas", "svg",
}
# 需要跳过的特定 class（搜索结果占位、评论区等噪音容器）
_SKIP_CLASSES = {"has-results", "no-results", "search-results"}

# 标题标签 → markdown 前缀
_HEADING_PREFIX = {
    "h1": "# ", "h2": "## ", "h3": "### ",
    "h4": "#### ", "h5": "##### ", "h6": "###### ",
}


def _table_to_markdown(table_node: lxml_html.HtmlElement) -> list[str]:
    """将 HTML <table> 转为 Markdown 表格，保留参数表结构。"""
    rows = table_node.xpath(".//tr")
    if not rows:
        return []

    # 分离表头和数据行
    header: list[str] = []
    data_rows: list[list[str]] = []
    for row in rows:
        cells = [c.text_content().strip() for c in row.xpath(".//th|.//td")]
        if not cells:
            continue
        is_header = bool(row.xpath(".//th"))
        if is_header or not header:
            # 使用第一行或显式 thead 行作为表头
            if not header:
                header = cells
            else:
                data_rows.append(cells)
        else:
            data_rows.append(cells)

    if not data_rows:
        return []

    lines: list[str] = ["", "| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for dr in data_rows:
        # 补齐列数
        while len(dr) < len(header):
            dr.append("")
        lines.append("| " + " | ".join(dr[:len(header)]) + " |")
    lines.append("")
    return lines


def _node_to_text(node: lxml_html.HtmlElement, depth: int = 0) -> list[str]:
    """递归遍历 DOM 节点，转为带标题层级的文本行列表。"""
    if node is None:
        return []

    tag = node.tag if isinstance(node.tag, str) else ""

    if tag in _SKIP_TAGS:
        return []

    # 跳过带噪音 class 的容器（搜索框、评论区等）
    cls = (node.get("class") or "").lower()
    if cls in _SKIP_CLASSES:
        return []

    # ── 行内元素：文本直接返回，由父级 _collect_inline_text 拼入 ──
    if tag in ("a", "strong", "b", "em", "i", "code", "span",
               "abbr", "sub", "sup", "small", "mark"):
        return []  # 文本已在父级的 _collect_inline_text 中处理

    # ── 收集子节点 ──
    child_block_lines: list[str] = []
    for child in node:
        child_output = _node_to_text(child, depth + 1)
        if child_output:
            ctag = child.tag if isinstance(child.tag, str) else ""
            if ctag in _BLOCK_TAGS | set(_HEADING_PREFIX.keys()) | {"pre", "li", "ul", "ol"}:
                if child_block_lines and child_block_lines[-1] != "":
                    child_block_lines.append("")
                child_block_lines.extend(child_output)
            # 行内子元素：文本已由 _collect_inline_text 收集，忽略

    lines: list[str] = []

    # ── 按标签类型输出 ──
    if tag in _HEADING_PREFIX:
        heading = _collect_inline_text(node)
        if heading:
            lines.append("")
            lines.append(_HEADING_PREFIX[tag] + heading)
            lines.append("")
        lines.extend(child_block_lines)

    elif tag == "table":
        lines.extend(_table_to_markdown(node))

    elif tag == "pre":
        # pre 内文本直接取
        code_text = node.text_content().strip() if hasattr(node, "text_content") else _collect_inline_text(node)
        if code_text:
            lines.append("")
            lines.append("```")
            lines.append(code_text)
            lines.append("```")
            lines.append("")

    elif tag == "li":
        # 判断父级是否为 <ol>，是则用数字序号
        parent_tag = node.getparent().tag if node.getparent() is not None else ""
        if parent_tag == "ol":
            # 计算在 ol 中的位置（只算前面的 li 兄弟节点）
            idx = 1
            prev = node.getprevious()
            while prev is not None:
                if prev.tag == "li":
                    idx += 1
                prev = prev.getprevious()
            prefix = f"{idx}. "
        else:
            prefix = "- "
        text = _collect_inline_text(node)
        if text:
            lines.append(prefix + text)
        else:
            # 没有行内文本时（li 里嵌套了 div/p），只输出子块内容
            pass
        lines.extend(child_block_lines)

    elif tag in _BLOCK_TAGS:
        block = _collect_inline_text(node)
        if block:
            lines.append(block)
        lines.extend(child_block_lines)

    elif tag == "img":
        alt = node.get("alt", "")
        if alt:
            lines.append(f"[图: {alt}]")

    elif tag == "br":
        lines.append("")

    else:
        other = _collect_inline_text(node)
        if other:
            lines.append(other)
        lines.extend(child_block_lines)

    return lines


def _collect_inline_text(node: lxml_html.HtmlElement) -> str:
    """收集节点及其行内子元素的文本，按文档顺序拼接。"""
    if node is None:
        return ""
    tag = node.tag if isinstance(node.tag, str) else ""
    if tag in _SKIP_TAGS:
        return ""

    parts: list[str] = []
    if node.text and node.text.strip():
        parts.append(node.text.strip())

    for child in node:
        ctag = child.tag if isinstance(child.tag, str) else ""
        if ctag in _SKIP_TAGS:
            continue
        if ctag in ("a", "strong", "b", "em", "i", "code", "span",
                     "abbr", "sub", "sup", "small", "mark"):
            child_text = _collect_inline_text(child)
            if child_text:
                parts.append(child_text)
        elif ctag == "img":
            alt = child.get("alt", "")
            if alt:
                parts.append(f"[图: {alt}]")
        elif ctag == "br":
            parts.append("\n")
        if child.tail and child.tail.strip():
            parts.append(child.tail.strip())

    return " ".join(parts)


def _postprocess(lines: list[str]) -> str:
    """后处理：合并多余空行、清理噪音标记。"""
    result: list[str] = []
    prev_empty = False
    for line in lines:
        stripped = line.strip()
        # 清理文档生成器的噪音标记
        if stripped in ("[TOC]", "[toc]", "[[toc]]", "{:toc}"):
            continue
        # 清理 GitBook 搜索框残留（"results matching" / "No results matching"）
        if stripped.startswith("# ") and ("results matching" in stripped.lower() or "no results" in stripped.lower()):
            continue
        if not stripped:
            if not prev_empty:
                result.append("")
                prev_empty = True
        else:
            result.append(stripped)
            prev_empty = False
    # 去掉首尾空行
    while result and result[0] == "":
        result.pop(0)
    while result and result[-1] == "":
        result.pop()
    return "\n".join(result)


def _try_extract_body(doc: lxml_html.HtmlElement) -> str:
    """尝试用 XPath 选择器定位正文区域，保留标题层级提取。

    按 _CONTENT_SELECTORS 优先级尝试，找到第一个非空区域即返回；
    全部未命中则回退到 <body> 全文。
    """
    for selector in _CONTENT_SELECTORS:
        nodes = doc.xpath(selector)
        if nodes:
            lines = _node_to_text(nodes[0])
            text = _postprocess(lines)
            if len(text) > 30:
                return text

    # 回退：取 <body> 全文
    body = doc.xpath("//body")
    if body:
        lines = _node_to_text(body[0])
        return _postprocess(lines)
    return ""


def _html_to_text(html_str: str) -> str:
    """HTML → 纯文本：清洗 → 正文定位 → 文本提取。"""
    doc = _parse_html_doc(html_str)
    doc = _clean_html(doc)
    return _try_extract_body(doc)


def _extract_title(html_str: str) -> str:
    """从 HTML 提取标题：XPath 查 <title>，取不到返回空串。"""
    try:
        doc = _parse_html_doc(html_str)
        titles = doc.xpath("//title/text()")
        return titles[0].strip() if titles else ""
    except Exception:
        # lxml 解析可能对畸形 HTML 抛异常，回退到字符串匹配
        lower = html_str.lower()
        if "<title>" in lower:
            start = lower.index("<title>") + 7
            end_tag = lower.index("</title>", start) if "</title>" in lower[start:] else len(lower)
            return html_str[start:end_tag].strip()
        return ""


def _extract_links(html_str: str) -> list[str]:
    """收集 HTML 中所有 <a href> 链接（XPath 实现，速度远快于 HTMLParser）。"""
    try:
        doc = _parse_html_doc(html_str)
        return [href for href in doc.xpath("//a/@href") if href]
    except Exception:
        return []


# 明显不是文档页面的扩展名，爬取时跳过
_SKIP_EXTS = (
    ".pdf", ".zip", ".png", ".jpg", ".jpeg", ".gif", ".svg", ".ico",
    ".css", ".js", ".json", ".xml", ".mp4", ".mp3", ".woff", ".woff2", ".ttf",
)


def _clean_pdf_text(text: str) -> str:
    """清洗 PDF 提取文本中的特殊字符。"""
    import unicodedata

    result: list[str] = []
    for ch in text:
        cat = unicodedata.category(ch)
        # 私有区字符、未分配字符 → 替换为普通空格
        if cat in ("Co", "Cn", "Cs"):
            result.append(" ")
        # PDF 中常见的特殊空格、零宽字符 → 普通空格或跳过
        elif ch in ("\xa0", "\xad", "\t", "​", "‌", "‍", "﻿"):
            result.append(" " if ch in ("\xa0", "\xad", "\t") else "")
        # 控制字符（除换行符）→ 跳过
        elif cat == "Cc" and ch not in ("\n", "\r"):
            continue
        else:
            result.append(ch)
    return "".join(result)


def _parse_pdf(data: bytes) -> str:
    """解析 PDF：用 pymupdf 提取文本、清洗、给编号标题加 ## 前缀。

    pymupdf 的文本提取质量远优于 pypdf，不会产生康熙部首乱码字符
    （如 ⼾→户、⼝→口）。但会保留原始排版换行，需后处理合并。
    """
    import pymupdf
    import re as _re

    doc = pymupdf.open(stream=data, filetype="pdf")
    raw_text = "\n".join(page.get_text() for page in doc)
    doc.close()

    raw_text = _clean_pdf_text(raw_text)

    # ── 合并被换行切断的短行（API 路径后缀、表头行） ──
    # 例如 "POST /v3/manage/query_tenant\n1" → "POST /v3/manage/query_tenant1"
    #      "create_resource\n1" → "create_resource1"
    lines = raw_text.split("\n")
    merged: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # 当前行非空，下一行是纯数字/短标识符（1-4字符），合并
        if line.strip() and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if nxt and _re.match(r"^\d{1,4}$", nxt) and len(nxt) <= 4:
                merged.append(line.rstrip() + nxt)
                i += 2
                continue
        merged.append(line)
        i += 1

    # 给编号标题行加 markdown 层级前缀
    # 1. / 2. → ##（H2），1.1 / 5.5 → ###（H3），1.1.1 → ####（H4）
    result = []
    for line in merged:
        stripped = line.strip()
        # 过滤纯数字行（页码、代码行号）
        if _re.match(r"^\d{1,6}$", stripped):
            continue
        # 过滤 PDF 乱码行（全小写无空格长串，如图片元数据）
        if _re.match(r"^[a-z]{10,}$", stripped):
            continue
        # 编号标题检测：覆盖有空格/无空格两种情况
        m = (
            _re.match(r"^(\d{1,2}(?:\.\d{1,2})+)\s+", stripped)       # 5.5 xxx
            or _re.match(r"^(\d{1,2}(?:\.\d{1,2})+)[一-鿿]", stripped) # 5.5中文(无空格)
            or _re.match(r"^(\d{1,2})\.[\s一-鿿]", stripped)           # 1.中文 或 1. 中文
            or _re.match(r"^(\d{1,2})\s+[一-鿿]", stripped)            # 1 中文
        )
        if m:
            if len(stripped) > 5:  # 至少有标题文字
                depth = m.group(1).count(".") + 1
                prefix = "#" * min(depth + 1, 4) + " "
                result.append(prefix + stripped)
                continue
        result.append(line)

    # ── 后处理：过滤连续数字行 + 代码块围栏 ──
    result = _filter_digit_runs(result)
    result = _fence_code_blocks(result)

    return "\n".join(result)


def _filter_digit_runs(lines: list[str]) -> list[str]:
    """过滤连续 3+ 个纯数字行（PDF 页码/代码行号）。"""
    out = list(lines)
    i = 0
    while i < len(out):
        if out[i].strip() and out[i].strip().isdigit():
            j = i
            while j < len(out) and out[j].strip().isdigit():
                j += 1
            if j - i >= 3:
                for k in range(i, j):
                    out[k] = ""
            i = j
        else:
            i += 1
    # 保留空行结构，只去掉被标记的数字行
    return out


def _fence_code_blocks(lines: list[str]) -> list[str]:
    """检测 JSON/代码块并加 ``` 围栏。"""
    out: list[str] = []
    in_block = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        starts_json = stripped.startswith("{") or stripped.startswith('["')

        if not in_block and starts_json:
            # 检查后续是否有延续（确认是代码块不是单独一行）
            continuation = 0
            for j in range(i + 1, min(i + 5, len(lines))):
                s = lines[j].strip()
                if s.startswith(('"', ',', '}', ']')) or s == "":
                    continuation += 1
                else:
                    break
            if continuation >= 2:
                in_block = True
                out.append("```")

        out.append(line)

        if in_block and (stripped.startswith("}") or stripped.startswith("]")):
            # 检查下一行是否还在块内
            next_stripped = lines[i + 1].strip() if i + 1 < len(lines) else ""
            if not next_stripped.startswith(('"', ',')):
                out.append("```")
                in_block = False

    if in_block:
        out.append("```")

    return out


def _parse_docx(data: bytes) -> str:
    """解析 DOCX，保留标题层级（Heading 1-6 → # ~ ######）。

    python-docx 可以读取段落的样式名，据此判断是否为标题。
    """
    from docx import Document
    import re as _re

    doc = Document(io.BytesIO(data))
    lines: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            lines.append("")  # 空行保留段落间距
            continue

        style = (p.style.name if p.style else "").lower()

        # 匹配 Word 内置标题样式：Heading 1, Heading 2, ...
        m = _re.match(r"heading\s*(\d+)", style)
        if m:
            level = int(m.group(1))
            if 1 <= level <= 6:
                prefix = "#" * level + " "
                lines.append(prefix + text)
                continue

        # 列表项样式
        if "list" in style or "bullet" in style:
            lines.append("- " + text)
            continue

        lines.append(text)

    # 表格内容
    for table in doc.tables:
        lines.append("")  # 表格前空行
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines).strip()


def _parse_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"# 工作表：{ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines).strip()


def parse_file(data: bytes, ext: str) -> str:
    """按扩展名把文件字节解析成文本。ext 形如 '.pdf'（小写、带点）。"""
    ext = ext.lower()
    if ext in PLAIN_EXTS:
        return data.decode("utf-8", errors="replace").strip()
    if ext in HTML_EXTS:
        return _html_to_text(data.decode("utf-8", errors="replace"))
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".docx":
        return _parse_docx(data)
    if ext == ".xlsx":
        return _parse_xlsx(data)
    raise ValueError(f"暂不支持的文件类型：{ext}")


def fetch_url(url: str) -> tuple[str, str, str]:
    """抓取网页，返回 (标题, 正文文本, 原始HTML)。"""
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    resp = httpx.get(
        url, headers={"User-Agent": _UA}, follow_redirects=True, timeout=20.0
    )
    resp.raise_for_status()
    html = resp.text
    text = _html_to_text(html)
    title = _extract_title(html) or url
    return title, text.strip(), html


def _norm(url: str) -> str:
    """规范化 URL 用于去重：去掉 #fragment 和结尾斜杠。"""
    clean, _ = urldefrag(url)
    return clean.rstrip("/")


def crawl_site(
    base_url: str, max_depth: int = 2, max_pages: int = 50
) -> tuple[str, str, int]:
    """从 base_url 递归抓取同域名、同路径前缀下的多个页面（BFS）。

    适合 GitBook/Docusaurus 等多层级文档站：一次把整套文档抓下来拼成一篇。

    Args:
        base_url: 起始 URL（通常是文档首页）
        max_depth: 最大链接深度（0=只抓起始页），默认 2
        max_pages: 最多抓多少页（安全上限），默认 50

    Returns:
        (整站标题, 拼接后的全文, 实际抓取页数)
    """
    if not base_url.startswith(("http://", "https://")):
        base_url = "https://" + base_url

    parsed = urlparse(base_url)
    domain = parsed.netloc
    # 路径前缀：取起始 URL 所在“目录”，限制只抓这个前缀下的页面
    # 例如 https://x.com/docs/intro -> 前缀 /docs/
    path = parsed.path
    if "/" in path:
        prefix = path.rsplit("/", 1)[0] + "/"
    else:
        prefix = "/"

    visited: set[str] = set()
    site_title = ""
    pages: list[tuple[str, str]] = []  # (页标题, 正文)

    # BFS 队列：(url, depth)
    queue: deque[tuple[str, int]] = deque([(base_url, 0)])
    visited.add(_norm(base_url))

    with httpx.Client(
        headers={"User-Agent": _UA}, follow_redirects=True, timeout=20.0
    ) as client:
        while queue and len(pages) < max_pages:
            url, depth = queue.popleft()
            try:
                resp = client.get(url)
                resp.raise_for_status()
            except Exception:
                continue  # 单页失败跳过，不中断整体
            ctype = resp.headers.get("content-type", "")
            if "html" not in ctype.lower():
                continue

            html = resp.text
            title = _extract_title(html) or url
            text = _html_to_text(html).strip()
            if text:
                pages.append((title, text))
            if not site_title:
                site_title = title  # 首页标题作整站标题

            if depth >= max_depth:
                continue

            # 提取站内链接，加入队列
            for href in _extract_links(html):
                if href.startswith(("mailto:", "javascript:", "tel:", "#")):
                    continue
                absu = urljoin(url, href)
                p = urlparse(absu)
                if p.scheme not in ("http", "https"):
                    continue
                if p.netloc != domain:  # 跨域不抓
                    continue
                if not p.path.startswith(prefix):  # 超出路径前缀不抓
                    continue
                if p.path.lower().endswith(_SKIP_EXTS):  # 非文档资源不抓
                    continue
                key = _norm(absu)
                if key in visited:
                    continue
                visited.add(key)
                queue.append((absu, depth + 1))

    if not pages:
        raise ValueError("未抓取到任何页面内容")

    # 拼接：每页用标题作分隔，方便 RAG 分块时保留章节边界
    combined = "\n\n".join(f"# {t}\n{body}" for t, body in pages)
    return site_title or base_url, combined, len(pages)
